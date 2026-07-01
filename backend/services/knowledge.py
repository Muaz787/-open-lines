import os
import asyncio
import logging
from datetime import datetime, timezone
from dotenv import load_dotenv
from firecrawl import FirecrawlApp
from openai import AsyncOpenAI
from pinecone import Pinecone

load_dotenv()

logger = logging.getLogger(__name__)

FIRECRAWL_API_KEY = os.getenv("FIRECRAWL_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")
PINECONE_INDEX_HOST = os.getenv("PINECONE_INDEX_HOST")

CHUNK_TOKENS = 500
OVERLAP_TOKENS = 50
# Approximation: 1 token ≈ 4 characters
CHARS_PER_TOKEN = 4

_openai: AsyncOpenAI | None = None
_pinecone_index = None


def _get_openai() -> AsyncOpenAI:
    global _openai
    if _openai is None:
        if not OPENAI_API_KEY:
            raise RuntimeError("OPENAI_API_KEY must be set")
        _openai = AsyncOpenAI(api_key=OPENAI_API_KEY)
    return _openai


def _get_pinecone_index():
    global _pinecone_index
    if _pinecone_index is None:
        if not PINECONE_API_KEY:
            raise RuntimeError("PINECONE_API_KEY must be set")
        if not PINECONE_INDEX_NAME and not PINECONE_INDEX_HOST:
            raise RuntimeError("PINECONE_INDEX_NAME or PINECONE_INDEX_HOST must be set")
        pc = Pinecone(api_key=PINECONE_API_KEY)
        if PINECONE_INDEX_HOST:
            # Strip scheme — pinecone-client 3.x adds https:// internally
            host = PINECONE_INDEX_HOST.removeprefix("https://").removeprefix("http://")
            _pinecone_index = pc.Index(host=host)
        else:
            _pinecone_index = pc.Index(PINECONE_INDEX_NAME)
    return _pinecone_index


def _chunk_text(text: str) -> list[str]:
    chunk_chars = CHUNK_TOKENS * CHARS_PER_TOKEN
    overlap_chars = OVERLAP_TOKENS * CHARS_PER_TOKEN
    step = chunk_chars - overlap_chars
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start : start + chunk_chars])
        start += step
    return chunks


async def scrape_website(url: str) -> str:
    if not FIRECRAWL_API_KEY:
        raise RuntimeError("FIRECRAWL_API_KEY must be set")
    app = FirecrawlApp(api_key=FIRECRAWL_API_KEY)
    result = app.scrape_url(url, params={"formats": ["markdown"]})
    # SDK may return a dict or an object with attribute access
    if isinstance(result, dict):
        text = result.get("markdown") or result.get("content", "")
    else:
        text = getattr(result, "markdown", None) or getattr(result, "content", "") or ""
    logger.info("Scraped %d chars from %s", len(text), url)
    return text


# Pages we never want in the knowledge base (legal/auth/cart/noise).
CRAWL_EXCLUDES = [
    "privacy", "terms", "cookie", "wp-admin", "wp-login", "login", "signin",
    "cart", "checkout", "account", "tag/", "/tag", "author/", "feed",
]
MAX_CRAWL_PAGES = 20


def _page_text(page) -> str:
    if isinstance(page, dict):
        return page.get("markdown") or page.get("content") or ""
    return getattr(page, "markdown", None) or getattr(page, "content", "") or ""


def _page_meta(page) -> dict:
    md = page.get("metadata") if isinstance(page, dict) else getattr(page, "metadata", None)
    md = md or {}
    get = (lambda k: md.get(k)) if isinstance(md, dict) else (lambda k: getattr(md, k, None))
    return {
        "source_url": str(get("sourceURL") or get("url") or "")[:400],
        "title": str(get("title") or "")[:200],
    }


async def crawl_website(url: str, max_pages: int = MAX_CRAWL_PAGES) -> tuple[str, list[dict]]:
    """Bounded, same-domain multi-page crawl via Firecrawl. Returns (combined
    markdown with per-page source headers, list of page metadata).

    Always safe: on ANY failure, empty result, or timeout it falls back to a
    single-page homepage scrape, so ingestion never breaks. Run off the event
    loop because the Firecrawl SDK is synchronous."""
    if not FIRECRAWL_API_KEY:
        raise RuntimeError("FIRECRAWL_API_KEY must be set")

    def _do_crawl() -> list:
        app = FirecrawlApp(api_key=FIRECRAWL_API_KEY)
        params = {
            "crawlerOptions": {
                "limit": max_pages,
                "maxDepth": 3,
                "excludes": CRAWL_EXCLUDES,
                "allowBackwardCrawling": False,
            },
            "pageOptions": {"onlyMainContent": True, "formats": ["markdown"]},
        }
        res = app.crawl_url(url, params=params, wait_until_done=True, timeout=120)
        # Normalise: SDK may return a list, or a dict with 'data'/'pages'.
        if isinstance(res, list):
            return res
        if isinstance(res, dict):
            return res.get("data") or res.get("pages") or []
        return getattr(res, "data", None) or []

    try:
        pages = await asyncio.to_thread(_do_crawl)
    except Exception as e:
        logger.warning("crawl_website: crawl failed for %s (%s) — falling back to single page", url, e)
        pages = []

    if not pages:
        text = await scrape_website(url)
        return text, [{"source_url": url, "title": ""}]

    blocks: list[str] = []
    meta: list[dict] = []
    for page in pages[:max_pages]:
        body = _page_text(page).strip()
        if not body:
            continue
        m = _page_meta(page)
        if not m["source_url"]:
            m["source_url"] = url
        meta.append(m)
        header = f"## Source: {m['source_url']}" + (f" — {m['title']}" if m["title"] else "")
        blocks.append(f"{header}\n\n{body}")

    if not blocks:  # crawl returned pages but no usable text
        text = await scrape_website(url)
        return text, [{"source_url": url, "title": ""}]

    combined = "\n\n---\n\n".join(blocks)
    logger.info("crawl_website: %s → %d pages, %d chars", url, len(meta), len(combined))
    return combined, meta


async def embed_and_store(
    namespace: str,
    text: str,
    tenant_id: str,
    source_id: str | None = None,
    source_type: str | None = None,
    extra_metadata: dict | None = None,
) -> int:
    """Embed `text` and upsert into Pinecone. `source_type` ('website' | 'file' |
    'text') tags every vector so website content can later be cleared without
    touching uploaded documents. `extra_metadata` (e.g. crawl_generation_id,
    crawled_at, page_count) is merged into every vector's metadata."""
    chunks = _chunk_text(text)
    if not chunks:
        logger.warning("No chunks produced for tenant %s", tenant_id)
        return 0

    openai = _get_openai()
    index = _get_pinecone_index()

    response = await openai.embeddings.create(
        model="text-embedding-3-small",
        input=chunks,
        dimensions=1024,
    )
    embeddings = [item.embedding for item in response.data]

    vectors = [
        {
            "id": f"{source_id}-{i}" if source_id else f"{tenant_id}-{i}",
            "values": embeddings[i],
            "metadata": {
                "text": chunks[i],
                "tenant_id": tenant_id,
                **({"source_id": source_id} if source_id else {}),
                **({"source_type": source_type} if source_type else {}),
                **(extra_metadata or {}),
            },
        }
        for i in range(len(chunks))
    ]

    batch_size = 100
    for start in range(0, len(vectors), batch_size):
        index.upsert(vectors=vectors[start : start + batch_size], namespace=namespace)

    logger.info(
        "Stored %d vectors in Pinecone namespace '%s' for tenant %s (source=%s, type=%s)",
        len(vectors), namespace, tenant_id, source_id or "none", source_type or "none",
    )
    return len(vectors)


def delete_by_source(namespace: str, source_id: str) -> None:
    index = _get_pinecone_index()
    index.delete(filter={"source_id": source_id}, namespace=namespace)
    logger.info("Deleted vectors for source %s in namespace '%s'", source_id, namespace)


def delete_old_website_vectors(namespace: str, keep_source_id: str) -> None:
    """Delete website vectors from EARLIER crawl generations, keeping only the
    just-written generation (`keep_source_id`). Never touches file/text vectors
    (different source_type). Safe failure mode: leaves duplicates, never empties."""
    index = _get_pinecone_index()
    index.delete(
        filter={"source_type": "website", "source_id": {"$ne": keep_source_id}},
        namespace=namespace,
    )


def delete_legacy_website_vectors(namespace: str, slug_or_tenant_ids: list[str]) -> None:
    """One-time cleanup of pre-tagging website vectors. Those used ids of the form
    `{tenant_id}-{i}` / `{slug}-{i}` with NO source_type, so they can't be reached
    by a metadata filter. Uploaded-document vectors use a UUID source_id prefix, so
    deleting these exact id ranges can never hit a document vector."""
    index = _get_pinecone_index()
    ids: list[str] = []
    for base in slug_or_tenant_ids:
        if base:
            ids.extend(f"{base}-{i}" for i in range(0, 1500))
    for start in range(0, len(ids), 1000):
        try:
            index.delete(ids=ids[start : start + 1000], namespace=namespace)
        except Exception as e:
            logger.warning("Legacy website vector cleanup batch failed for '%s': %s", namespace, e)
            return


IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp")
CHARS_PER_PAGE_EST = 2500  # rough per-page text budget for non-paged formats


async def extract_document(
    filename: str, content: bytes, *, max_pages: int, ocr_pages: int,
) -> dict:
    """Extract text with per-plan page/OCR caps. Returns
    {text, pages, truncated, ocr_used}. Scanned PDFs and images route to Mistral OCR."""
    from services import ocr
    name = (filename or "").lower()
    truncated = False
    ocr_used = False
    pages_total: int | None = None

    if name.endswith(".pdf"):
        text, pages_total, truncated = _extract_pdf_capped(content, max_pages)
        if len(text.strip()) < 40:  # no usable text layer → likely scanned → OCR
            ocr_text = await ocr.ocr_document(content, filename, max_pages=ocr_pages)
            if ocr_text.strip():
                text = ocr_text
                ocr_used = True
                truncated = bool(pages_total and pages_total > ocr_pages)
    elif name.endswith(IMAGE_EXTS):
        text = await ocr.ocr_document(content, filename, max_pages=1)
        ocr_used = True
        pages_total = 1
    elif name.endswith(".docx"):
        text = _extract_docx(content)
    elif name.endswith(".xlsx"):
        text = _extract_xlsx(content)
    elif name.endswith(".xls"):
        text = _extract_xls(content)
    elif name.endswith((".txt", ".md", ".csv")):
        text = content.decode("utf-8", errors="ignore")
    else:
        ext = name.rsplit(".", 1)[-1] if "." in name else "unknown"
        raise ValueError(
            f"Unsupported file type .{ext}. Accepted: PDF, Word (.docx), "
            f"Excel (.xlsx/.xls), images (.png/.jpg), plain text (.txt/.csv/.md)"
        )

    # Char safety cap for non-paged / very long text (scales with the plan's page cap).
    char_cap = max_pages * CHARS_PER_PAGE_EST
    if len(text) > char_cap:
        text = text[:char_cap]
        truncated = True

    return {"text": text, "pages": pages_total, "truncated": truncated, "ocr_used": ocr_used}


def extract_text(filename: str, content: bytes) -> str:
    """Sync text extraction (no OCR, no page cap) — kept for non-upload callers."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    elif name.endswith(".docx"):
        return _extract_docx(content)
    elif name.endswith(".xlsx"):
        return _extract_xlsx(content)
    elif name.endswith(".xls"):
        return _extract_xls(content)
    elif name.endswith((".txt", ".md", ".csv")):
        return content.decode("utf-8", errors="ignore")
    else:
        ext = name.rsplit(".", 1)[-1] if "." in name else "unknown"
        raise ValueError(f"Unsupported file type .{ext}. Accepted: PDF, Word (.docx), Excel (.xlsx/.xls), plain text (.txt/.csv/.md)")


def _extract_pdf(content: bytes) -> str:
    import io
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _extract_pdf_capped(content: bytes, max_pages: int) -> tuple[str, int, bool]:
    """Extract the text layer from the first `max_pages` pages.
    Returns (text, total_pages, truncated)."""
    import io
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    total = len(reader.pages)
    use = min(total, max(1, max_pages))
    pages = [reader.pages[i].extract_text() or "" for i in range(use)]
    return "\n\n".join(p for p in pages if p.strip()), total, total > max_pages


def _extract_xls(content: bytes) -> str:
    """Legacy .xls (BIFF) via xlrd — openpyxl only reads .xlsx."""
    import xlrd
    book = xlrd.open_workbook(file_contents=content)
    rows: list[str] = []
    for sheet in book.sheets():
        rows.append(f"[Sheet: {sheet.name}]")
        for r in range(sheet.nrows):
            cells = [str(sheet.cell_value(r, ci)).strip() for ci in range(sheet.ncols)]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
    return "\n".join(rows)


def _extract_docx(content: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    import io
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    rows: list[str] = []
    for sheet in wb.worksheets:
        rows.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))
    wb.close()
    return "\n".join(rows)


async def query_knowledge_base(namespace: str, query: str, top_k: int = 5) -> str:
    openai = _get_openai()
    index = _get_pinecone_index()

    response = await openai.embeddings.create(
        model="text-embedding-3-small",
        input=[query],
        dimensions=1024,
    )
    query_vector = response.data[0].embedding

    results = index.query(
        vector=query_vector,
        top_k=top_k,
        namespace=namespace,
        include_metadata=True,
    )

    matches = results.get("matches", [])
    texts = [m["metadata"]["text"] for m in matches if m.get("metadata", {}).get("text")]
    return "\n\n".join(texts)


# Structured queries that pull the SPECIFIC facts a receptionist needs, instead of
# one generic "overview" query (which biased toward homepage copy and missed
# pricing / hours / policies). Captures website AND uploaded-document chunks.
_STRUCTURED_QUERIES: list[tuple[str, int]] = [
    ("services offered and what the business does", 6),
    ("pricing, fees, rates, and costs", 4),
    ("business hours, opening times, and availability", 3),
    ("service area, locations, and neighbourhoods served", 3),
    ("frequently asked questions and answers", 5),
    ("warranty, guarantee, cancellation, emergency, and payment policies", 4),
]


async def build_structured_knowledge(
    namespace: str, business_brief: str = "", max_chars: int = 12000
) -> str:
    """Assemble the factual knowledge baked into the assistant prompt: the stored
    Business Brief first, then de-duplicated chunks from several targeted Pinecone
    queries (so specifics like pricing, hours, and policies surface — and uploaded
    documents are included). Capped to keep the prompt lean (no call-time latency)."""
    parts: list[str] = []
    if business_brief and business_brief.strip():
        parts.append(business_brief.strip())

    seen: set[str] = set()
    for query, k in _STRUCTURED_QUERIES:
        try:
            text = await query_knowledge_base(namespace, query, top_k=k)
        except Exception as e:
            logger.warning("Structured KB query '%s' failed for '%s': %s", query, namespace, e)
            continue
        for chunk in text.split("\n\n"):
            c = chunk.strip()
            key = c[:120]
            if c and key not in seen:
                seen.add(key)
                parts.append(c)

    combined = "\n\n".join(parts).strip()
    return combined[:max_chars] if combined else ""


def clear_namespace(namespace: str) -> None:
    index = _get_pinecone_index()
    index.delete(delete_all=True, namespace=namespace)
    logger.info("Cleared Pinecone namespace '%s'", namespace)


MAX_SCRAPE_CHARS = 200_000  # cap embedding cost per crawl (scrape is single-page)


async def refresh_tenant_knowledge(
    tenant_id: str, website_url: str, namespace: str
) -> dict:
    """Two-phase, safe website refresh:
      1. Scrape first — if it fails or returns nothing, raise, so the caller keeps
         the existing knowledge base rather than emptying it.
      2. Embed the NEW crawl generation, THEN delete older website generations
         (and pre-tagging legacy website vectors). Cleanup runs only after the new
         content is in place, so a failure leaves duplicates — never an empty KB —
         and uploaded-document vectors are never touched.
    """
    import time

    # Phase 1 — bounded multi-page crawl (falls back to single page internally).
    # Most likely failure point, so it runs first: if nothing comes back we raise
    # and the caller keeps the existing KB.
    raw_text, pages_meta = await crawl_website(website_url, max_pages=MAX_CRAWL_PAGES)
    if not raw_text or not raw_text.strip():
        raise ValueError("Website crawl returned no content")
    raw_text = raw_text[:MAX_SCRAPE_CHARS]
    page_count = len(pages_meta) or (raw_text.count("\n\n") + 1)
    crawled_at = datetime.now(timezone.utc)

    # Regenerate the structured Business Brief from the full crawl (best-effort).
    brief_data: dict = {}
    try:
        from services import website_analysis
        brief_data = await website_analysis.classify_text(raw_text, website_url, max_chars=40000)
    except Exception as e:
        logger.warning("Brief regeneration failed for tenant %s (non-fatal): %s", tenant_id, e)

    # Phase 2 — write the new generation under a unique website source id.
    gen_source_id = f"website-{int(time.time())}"
    vectors_stored = await embed_and_store(
        namespace, raw_text, tenant_id,
        source_id=gen_source_id, source_type="website",
        extra_metadata={"crawl_generation_id": gen_source_id, "crawled_at": crawled_at.isoformat(), "page_count": page_count},
    )

    # Remove previous website generations + legacy untagged website vectors.
    try:
        delete_old_website_vectors(namespace, keep_source_id=gen_source_id)
        delete_legacy_website_vectors(namespace, list({namespace, tenant_id}))
    except Exception as e:
        logger.warning("Old website vector cleanup failed for tenant %s (non-fatal): %s", tenant_id, e)

    logger.info(
        "Knowledge refresh complete for tenant %s: %d pages, %d vectors",
        tenant_id, page_count, vectors_stored,
    )
    return {
        "pages_scraped": page_count,
        "vectors_stored": vectors_stored,
        "refreshed_at": crawled_at,
        "brief": brief_data,
        "pages_meta": pages_meta,
    }
